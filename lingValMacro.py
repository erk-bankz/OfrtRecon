# basic imports
from lingValUI import Ui_Dialog  # make baby
from PyQt5.QtCore import pyqtSlot
from PyQt5.QtWidgets import QDialog, QApplication
import openpyxl
import docx
from pathlib import Path
import re
from PyQt5 import QtWidgets
import modules


class Dialog(QDialog, Ui_Dialog):
    """
    Copy and paste from ORT to excel or vice versa
    """
    def __init__(self, parent=None):
        super(Dialog, self).__init__(parent)
        self.setupUi(self)
    #     self.pushButton.clicked.connect(self.buttonClicked)  # alternative way to call your method
    #
    # def buttonClicked(self):
    #     # write code

    @pyqtSlot()  # decorator
    def on_pushButton_clicked(self):
        """
        Write what this method does
        """

        ft_file_path = Path(self.lineEditPath1.text())
        bt_file_path = Path(self.lineEditPath2.text())
        xls_file_path = Path(self.lineEditPath3.text())
        modules.removeCheckbox(xls_file_path)
        if self.checkBox.isChecked():
            hideTwoRows = True
        else:
            hideTwoRows = False
        front_user_column = int(self.excelSourceColumn.text())
        back_user_column = int(self.excelTargetColumn.text())
        for file in ft_file_path.iterdir():
            if file.suffix == ".docx":
                front_doc, front_lp = modules.extractfileNameandFileLP(file)
                if front_lp == "":
                    QtWidgets.QMessageBox.information(self, "Error!", "Please make sure LP is present in the file name and in '-xx-XX' format ")
                    exit()
                bt_file_name = modules.findBTmatch(front_doc, front_lp, bt_file_path)
                xls_file_name = modules.findXLSmatch(front_doc, front_lp, xls_file_path)
                if bt_file_name is None:
                    QtWidgets.QMessageBox.information(self, "Error!", "BT file match not found, skipping {0}".format(file.name))
                    continue
                if xls_file_name is None:
                    print(front_doc)
                    print(front_lp)
                    QtWidgets.QMessageBox.information(self, "Error!", "XLS match not found, skipping {0}".format(file.name))
                    continue
                front_doc = docx.Document(file)
                back_doc = docx.Document(bt_file_name)
                xls_file = openpyxl.load_workbook(xls_file_name)
                ft_ort_table = modules.extract_table_values(front_doc.tables[2])
                bt_ort_table = modules.extract_table_values(back_doc.tables[2])
                active_sheet = xls_file.active
                if hideTwoRows:
                    active_sheet.row_dimensions[1].hidden = True
                    active_sheet.row_dimensions[2].hidden = True
                    xls_file.save(str(xls_file_name))
                if self.rbOption1.isChecked():
                    modules.RecontoORT(front_doc.tables[2], back_doc.tables[2], xls_file, front_user_column, back_user_column)
                    front_doc.save(str(file))
                    back_doc.save(str(bt_file_name))
                if self.rbOption2.isChecked():
                    modules.ORTtoRecon(ft_ort_table, bt_ort_table, active_sheet,front_user_column, back_user_column)
                    xls_file.save(str(xls_file_name))
            elif file.suffix == ".doc":
                QtWidgets.QMessageBox.information(self, "Wrong file format", "{0} is a doc, please convert to docx".format(file.name))
                continue
        QtWidgets.QMessageBox.information(self, "Done!", "Macro finished")



if __name__ == "__main__":
    import sys
    application = QApplication(sys.argv)
    macro_dialog = Dialog() # create object of dialog, **use the name of your dialog**
    macro_dialog.show()
    sys.exit(application.exec_())



